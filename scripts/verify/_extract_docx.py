"""docx 추출 — R-ONE 미분양 STATBL_ID 확인용."""
from pathlib import Path
import docx

DOC = Path(r"C:\Users\leejh\OneDrive\바탕 화면\2026년 국토교통 데이터 활용 경진대회\기술문서_부동산통계 Open API 서비스_240905.docx")
d = docx.Document(str(DOC))
print("===PARAGRAPHS===")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[P{i}] {t}")
print()
print("===TABLES===")
for ti, tbl in enumerate(d.tables):
    print(f"---TABLE {ti}---")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
        print(f"  R{ri}: " + " | ".join(cells))
